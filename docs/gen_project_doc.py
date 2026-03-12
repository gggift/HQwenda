"""
Generate HQwenda project documentation in GB/T 9704-2012 format.
"""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os


def set_cell_font(cell, font_name, font_size, bold=False):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            run.font.size = Pt(font_size)
            run.font.bold = bold


def create_doc():
    doc = Document()

    # ── Page layout: A4 with GB/T 9704-2012 margins ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # ── Helper: set run font ──
    def set_run(run, font_name, size_pt, bold=False):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(size_pt)
        run.font.bold = bold

    # ── Helper: set paragraph line spacing to 28pt fixed ──
    def set_line_spacing(paragraph, pt=28):
        pf = paragraph.paragraph_format
        pf.line_spacing = Pt(pt)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    # ── Helper: add title (二号小标宋, centered) ──
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_line_spacing(p, 28)
        run = p.add_run(text)
        set_run(run, "华文中宋", 22, bold=False)  # 二号 ≈ 22pt, 小标宋→华文中宋

    # ── Helper: add level-1 heading (三号黑体, "一、XXX") ──
    def add_h1(text):
        p = doc.add_paragraph()
        set_line_spacing(p, 28)
        pf = p.paragraph_format
        pf.first_line_indent = Emu(0)
        pf.space_before = Pt(6)
        run = p.add_run(text)
        set_run(run, "黑体", 16, bold=False)  # 三号 ≈ 16pt

    # ── Helper: add level-2 heading (三号楷体加粗, "（一）XXX") ──
    def add_h2(text):
        p = doc.add_paragraph()
        set_line_spacing(p, 28)
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0.74)  # 2字符
        run = p.add_run(text)
        set_run(run, "楷体", 16, bold=True)

    # ── Helper: add level-3 heading (三号仿宋加粗, "1. XXX") ──
    def add_h3(text):
        p = doc.add_paragraph()
        set_line_spacing(p, 28)
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0.74)
        run = p.add_run(text)
        set_run(run, "仿宋", 16, bold=True)

    # ── Helper: add body paragraph (三号仿宋, indent 2 chars) ──
    def add_body(text):
        p = doc.add_paragraph()
        set_line_spacing(p, 28)
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0.74)  # 2字符 ≈ 0.74cm at 三号
        run = p.add_run(text)
        set_run(run, "仿宋", 16)

    # ── Helper: add code block (小四号等线) ──
    def add_code(text):
        p = doc.add_paragraph()
        set_line_spacing(p, 24)
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0.74)
        run = p.add_run(text)
        set_run(run, "Consolas", 12)

    # ══════════════════════════════════════════════
    #  DOCUMENT CONTENT
    # ══════════════════════════════════════════════

    # ── Title ──
    add_title("HQwenda金融行情数据智能问答系统")
    add_title("技术设计文档")

    # Blank line
    p = doc.add_paragraph()
    set_line_spacing(p, 28)

    # ── 一、项目概述 ──
    add_h1("一、项目概述")

    add_h2("（一）项目背景")
    add_body(
        "随着国内资本市场数据量持续增长，专业投资者和研究人员对金融数据的即时查询需求日益迫切。"
        "传统的数据终端操作复杂、学习成本高，难以满足快速问答场景。"
        "HQwenda系统旨在构建一个基于大语言模型（LLM）Function Calling能力的金融数据智能问答平台，"
        "使用户能够通过自然语言直接获取股票行情、指数走势、行业分析、资金流向等结构化金融数据。"
    )

    add_h2("（二）系统定位")
    add_body(
        "本系统定位为专业级金融数据问答助手，核心特点如下："
    )
    add_body("1. 自然语言交互：用户以中文自然语言提问，系统自动解析意图并调用相应数据接口。")
    add_body("2. 结构化数据驱动：所有回答均基于Tushare金融数据接口的实时和历史数据，确保数据准确性。")
    add_body("3. Agent工具调用架构：采用LLM Function Calling模式，由大模型自主决定调用哪些工具、传递什么参数。")
    add_body("4. 多轮会话支持：支持上下文连续对话，最多保持20轮历史记录。")

    add_h2("（三）技术选型")
    add_body("1. 后端框架：FastAPI（Python），同步线程池模式处理阻塞式API调用。")
    add_body("2. 大语言模型：DeepSeek V3（deepseek-chat-v3-0324），通过OpenRouter API路由调用。")
    add_body("3. 金融数据源：Tushare Pro API，覆盖A股、指数、行业、宏观、期货、外汇等数据。")
    add_body("4. 前端：原生HTML/CSS/JavaScript单页应用，轻量级部署。")
    add_body("5. 部署方式：Uvicorn ASGI服务器，systemd守护进程管理。")

    # ── 二、系统架构 ──
    add_h1("二、系统架构")

    add_h2("（一）整体架构")
    add_body(
        "系统采用前后端分离的单体架构，核心由以下五层组成："
    )
    add_body(
        "1. 表现层（Frontend）：HTML/CSS/JS单页应用，负责用户交互和消息渲染，"
        "支持Markdown表格、代码块等富文本展示。"
    )
    add_body(
        "2. 接口层（API Layer）：FastAPI路由，提供/api/chat（对话接口）、"
        "/api/session/new（创建会话）、/api/session/{id}（删除会话）等RESTful端点。"
    )
    add_body(
        "3. 智能体层（Agent Layer）：核心推理引擎，负责与LLM交互，"
        "执行多轮工具调用循环，直至生成最终回答。"
    )
    add_body(
        "4. 工具层（Tool Layer）：31个注册工具函数，覆盖行情、指数、行业、"
        "资金流向、技术分析、宏观数据等六大类金融数据查询。"
    )
    add_body(
        "5. 数据层（Data Layer）：Tushare Pro API提供底层金融数据，"
        "知识库（Markdown文件）提供辅助领域知识。"
    )

    add_h2("（二）模块结构")
    add_body("系统目录结构如下：")
    add_code("app/")
    add_code("├── main.py              # FastAPI应用入口")
    add_code("├── config.py            # 配置管理（环境变量）")
    add_code("├── api/chat.py          # 对话API端点")
    add_code("├── agent/")
    add_code("│   ├── core.py          # Agent推理循环")
    add_code("│   └── context.py       # 上下文组装")
    add_code("├── tools/")
    add_code("│   ├── registry.py      # 工具注册中心")
    add_code("│   ├── utils_tool.py    # 基础工具（股票查询、交易日历）")
    add_code("│   ├── quotes.py        # 行情工具（7个）")
    add_code("│   ├── index.py         # 指数工具（5个）")
    add_code("│   ├── fundamental.py   # 基本面工具（3个）")
    add_code("│   ├── industry.py      # 行业工具（4个）")
    add_code("│   ├── technical.py     # 技术分析工具（1个）")
    add_code("│   ├── moneyflow.py     # 资金流向工具（2个）")
    add_code("│   ├── macro.py         # 宏观数据工具（5个）")
    add_code("│   └── concept.py       # 概念板块工具（2个）")
    add_code("├── session/manager.py   # 会话管理")
    add_code("├── knowledge/           # 知识库（Markdown文件）")
    add_code("└── static/              # 前端静态文件")

    add_h2("（三）数据流向")
    add_body(
        "系统处理一次用户提问的完整数据流如下："
    )
    add_body(
        "1. 用户在前端输入自然语言问题，前端通过POST /api/chat发送请求。"
    )
    add_body(
        "2. 后端接口层接收请求，调用上下文组装模块（context.py），"
        "提取关键词匹配知识库文档，获取会话历史，构建系统提示词。"
    )
    add_body(
        "3. Agent核心（core.py）将系统提示词、历史消息和工具Schema列表一并发送给LLM。"
    )
    add_body(
        "4. LLM分析用户意图，返回需要调用的工具名称和参数（Function Calling）。"
    )
    add_body(
        "5. Agent执行工具调用，获取Tushare数据，将结果追加到消息列表。"
    )
    add_body(
        "6. Agent再次调用LLM，LLM根据工具返回的数据生成自然语言回答。"
    )
    add_body(
        "7. 若LLM判断还需更多数据，重复步骤4至6，最多循环10次。"
    )
    add_body(
        "8. 最终回答返回前端展示，同时将对话记录保存至会话管理器。"
    )

    # ── 三、问题拆解与处理流程 ──
    add_h1("三、问题拆解与处理流程")

    add_h2("（一）问题分类")
    add_body("系统支持的问题类型涵盖以下六大类：")
    add_body(
        "1. 行情查询类：个股日K线、周K线、月K线、实时行情、"
        "全市场涨跌统计（涨停/跌停数量）、全市场市值排名。"
    )
    add_body(
        "2. 指数分析类：境内指数日行情（上证、深证、创业板等）、"
        "境外指数（标普500、恒生、日经等）、指数成分股权重、"
        "成分股涨跌统计、历史条件筛选（如'过去10年收盘高于4000点的天数'）。"
    )
    add_body(
        "3. 行业研究类：申万一级行业涨跌排名、二级/三级行业数据、"
        "行业对指数贡献点数和涨跌幅、行业PE/PB估值排名、同花顺概念板块。"
    )
    add_body(
        "4. 资金流向类：北向资金（沪深港通）每日流入流出、"
        "个股资金流向（超大单/大单/中单/小单买卖金额）。"
    )
    add_body(
        "5. 基本面与估值类：个股PE/PB/PS/市值/换手率/股息率、"
        "财务指标（ROE/ROA/毛利率/净利率/EPS）、历史分位数计算。"
    )
    add_body(
        "6. 宏观与衍生品类：人民币汇率（USD/CNY等）、"
        "SHIBOR利率、中国国债收益率曲线、美国国债收益率、期货行情（黄金/原油/铜等）。"
    )

    add_h2("（二）处理流程详解")
    add_body(
        "以用户提问'昨天申万一级行业对上证指数的贡献点数排名'为例，"
        "系统处理流程如下："
    )
    add_h3("1. 上下文组装阶段")
    add_body(
        "系统从用户消息中提取关键词（'指数''上证'），"
        "匹配知识库中的相关文档。注入当前日期到系统提示词中，"
        "确保LLM知道'昨天'对应的具体日期。检索会话历史记录，"
        "组装完整的消息列表。"
    )
    add_h3("2. LLM意图识别阶段")
    add_body(
        "LLM接收系统提示词和31个工具的Schema描述，分析用户意图，"
        "决定调用get_industry_index_contribution工具，"
        "自动填充参数index_code为000001.SH、trade_date为昨天的日期。"
    )
    add_h3("3. 工具执行阶段")
    add_body(
        "工具函数执行以下步骤：通过index_member获取31个一级行业的成份股映射；"
        "通过daily_basic获取流通市值作为权重估算；"
        "通过daily获取全市场个股涨跌幅；"
        "按行业汇总加权贡献度；通过index_daily获取指数收盘价换算为贡献点数。"
    )
    add_h3("4. 回答生成阶段")
    add_body(
        "LLM收到工具返回的结构化数据后，"
        "将其组织为按贡献点数排序的列表，附带行业权重和涨跌幅信息，"
        "生成符合用户预期的中文回答。"
    )

    add_h2("（三）多轮工具调用机制")
    add_body(
        "系统采用循环调用机制，最多执行10轮LLM与工具的交互。"
        "典型的多轮场景包括：第一轮调用get_stock_basic将股票名称解析为代码，"
        "第二轮调用get_daily_quotes获取行情数据。"
        "每轮的工具调用结果均追加到消息历史中，供LLM在后续轮次参考。"
        "当LLM判断已获取足够信息时，直接输出文本回答，循环终止。"
    )

    # ── 四、提示词工程 ──
    add_h1("四、提示词工程")

    add_h2("（一）系统提示词设计")
    add_body("系统提示词由以下四个部分动态组装：")
    add_h3("1. 角色定义与日期注入")
    add_body(
        "提示词首先声明助手角色为'专业的金融数据分析助手'，"
        "并注入当前日期（同时提供YYYY-MM-DD和YYYYMMDD两种格式），"
        "确保LLM在处理'今天''昨天''本周'等相对时间表述时能够正确转换为绝对日期。"
    )
    add_h3("2. 行为约束")
    add_body(
        "明确指示LLM：用户会问关于股票、指数等金融数据的问题，"
        "可以调用工具获取实时数据来回答问题，回答要准确、简洁。"
        "当用户提到股票名称时，先用get_stock_basic工具查找股票代码，再用代码查询数据。"
    )
    add_h3("3. 知识库注入")
    add_body(
        "根据用户消息中的关键词（如PE、市盈率、均线、交易规则等），"
        "从Markdown知识库中检索相关文档，作为'相关知识'附加到系统提示词末尾，"
        "为LLM提供领域背景知识。"
    )
    add_h3("4. 完整提示词示例")
    add_body("以下为一次典型请求的系统提示词全文：")
    add_code("你是一个专业的金融数据分析助手。今天的日期是 2026-03-12")
    add_code("（工具调用时使用 20260312 格式）。")
    add_code("用户会问你关于股票、指数等金融数据的问题。")
    add_code("你可以调用工具获取实时数据来回答问题。回答要准确、简洁。")
    add_code("当用户提到股票名称时，先用 get_stock_basic 工具查找股票代码，")
    add_code("再用代码查询数据。")

    add_h2("（二）工具Schema设计原则")
    add_body(
        "每个工具通过JSON Schema定义其名称、描述和参数，供LLM理解工具用途。"
        "Schema设计遵循以下原则："
    )
    add_body(
        "1. 描述精确：工具描述中明确说明数据内容、单位、适用场景。"
        "例如get_sw_daily的描述为'获取申万行业指数日行情数据，"
        "支持按行业级别筛选（L1一级行业31个、L2二级行业、L3三级行业），默认只返回一级行业'。"
    )
    add_body(
        "2. 参数语义化：每个参数附带清晰的description，"
        "包含格式要求和示例值。如trade_date标注'格式YYYYMMDD'，"
        "ts_code标注'如000001.SZ'。"
    )
    add_body(
        "3. 返回值注释：工具返回的dict中附带note字段，"
        "解释各字段含义和单位，帮助LLM正确解读数据。"
    )

    # ── 五、工具集清单 ──
    add_h1("五、工具集清单")

    add_h2("（一）基础工具（3个）")
    add_body("1. get_stock_basic：获取股票基础信息，支持按名称或代码查询。")
    add_body("2. get_trade_calendar：获取交易日历，查询某段时间的交易日和休市日。")
    add_body("3. calculate_metric：简单计算工具，支持涨跌幅、均值、最大值、最小值、求和。")

    add_h2("（二）行情工具（7个）")
    add_body("1. get_daily_quotes：个股日K线数据，含成交额（亿元）、涨跌幅。")
    add_body("2. get_weekly_monthly：个股周K线和月K线数据。")
    add_body("3. get_market_stats：全市场涨跌统计，含涨停/跌停家数。")
    add_body("4. get_multi_period_returns：个股多周期累计涨跌幅（5日/20日/60日/120日/250日）。")
    add_body("5. get_industry_valuation_rank：申万一级行业PE/PB估值排名。")
    add_body("6. get_realtime_quote：个股实时行情（当前价、开盘价、最高最低价等）。")
    add_body("7. get_historical_percentile：历史分位数计算（价格/成交额/PE/PB等指标）。")

    add_h2("（三）指数工具（5个）")
    add_body("1. get_index_daily：境内指数日行情数据。")
    add_body("2. get_index_weight：指数成分股及权重。")
    add_body("3. get_index_history_filter：长周期历史数据条件筛选。")
    add_body("4. get_index_member_stats：指数成分股涨跌统计及权重贡献排名。")
    add_body("5. get_index_global：境外主要指数行情。")

    add_h2("（四）行业工具（4个）")
    add_body("1. get_sw_daily：申万行业指数日行情，支持L1/L2/L3级别筛选。")
    add_body("2. get_industry_index_contribution：行业对指数贡献点数和涨跌幅计算。")
    add_body("3. get_ths_daily：同花顺概念板块日行情。")
    add_body("4. get_market_rank：全市场按市值/PE/PB/换手率/股息率排名。")

    add_h2("（五）资金流向工具（2个）")
    add_body("1. get_northbound_flow：沪深港通北向资金每日流入流出数据。")
    add_body("2. get_moneyflow：个股资金流向（超大单/大单/中单/小单）。")

    add_h2("（六）技术分析工具（1个）")
    add_body("1. get_moving_averages：均线系统（MA5至MA250），含金叉/死叉检测和乖离率。")

    add_h2("（七）宏观与衍生品工具（5个）")
    add_body("1. get_fx_daily：外汇日行情（美元/人民币等）。")
    add_body("2. get_shibor：上海银行间同业拆放利率。")
    add_body("3. get_cn_bond_yield：中国国债收益率曲线。")
    add_body("4. get_us_treasury_yield：美国国债收益率（1月至30年期）。")
    add_body("5. get_fut_daily：期货日行情（黄金、原油、铜等）。")

    add_h2("（八）概念板块工具（2个）")
    add_body("1. get_concept_list：获取概念板块列表。")
    add_body("2. get_concept_stocks：获取概念板块成分股。")

    # ── 六、关键设计决策 ──
    add_h1("六、关键设计决策")

    add_h2("（一）数据单位标准化")
    add_body(
        "Tushare接口返回的成交额（amount）原始单位为千元，"
        "系统在工具层统一除以100000转换为亿元，"
        "并在工具描述和返回值note中标注'单位亿元'，"
        "避免LLM误读数量级导致回答错误。"
        "北向资金数据原始单位为百万元，统一转换为亿元。"
    )

    add_h2("（二）申万行业分级过滤")
    add_body(
        "Tushare的sw_daily接口返回所有级别行业数据（约439条，含一级、二级、三级行业），"
        "系统通过index_classify接口获取31个一级行业代码并缓存，"
        "在查询时默认过滤为一级行业，避免数据混淆。"
        "行业对指数贡献度计算通过index_member接口逐一获取31个一级行业的成份股列表，"
        "建立精确的个股到一级行业名称映射（缓存机制，首次加载后常驻内存）。"
    )

    add_h2("（三）日期感知机制")
    add_body(
        "系统在每次请求时将当前日期注入系统提示词，"
        "同时提供两种格式（YYYY-MM-DD和YYYYMMDD），"
        "确保LLM能正确处理'今天''昨天''本周'等相对时间表述。"
        "部分工具增加了日期回退机制：当指定日期无数据时（如当天交易未结束），"
        "自动向前查找最近的有效交易日。"
    )

    add_h2("（四）指数权重估算策略")
    add_body(
        "对于提供精确权重数据的指数（如沪深300、上证50），"
        "系统直接使用index_weight接口获取官方权重。"
        "对于全市场综合指数（如上证综指），"
        "由于无法获取逐只股票的精确权重，"
        "系统使用流通市值（circ_mv）按比例估算权重，"
        "并在返回结果中标注权重计算方法。"
    )

    add_h2("（五）错误容错与降级")
    add_body(
        "前端设置180秒超时，防止长时间无响应。"
        "后端Agent循环最多10轮，防止无限递归。"
        "所有工具函数均包含异常捕获，返回结构化错误信息而非抛出异常。"
        "API端点使用try/except包裹，确保任何内部错误均返回友好的中文提示。"
    )

    # ── 七、部署方案 ──
    add_h1("七、部署方案")

    add_h2("（一）服务器环境")
    add_body("操作系统：Ubuntu 22.04 LTS（Linux 5.15内核）。")
    add_body("Python版本：3.10.12。")
    add_body("服务器配置：阿里云ECS实例（公网IP：8.147.71.40）。")

    add_h2("（二）部署架构")
    add_body(
        "系统以systemd守护进程方式运行，服务名称为hqwenda，"
        "项目部署路径为/opt/HQwenda，"
        "通过Uvicorn ASGI服务器监听8000端口。"
        "服务配置为开机自启动，异常退出后5秒自动重启。"
    )

    add_h2("（三）环境变量配置")
    add_body("系统通过.env文件管理敏感配置，主要包括：")
    add_body("1. DEEPSEEK_API_KEY：OpenRouter API密钥。")
    add_body("2. DEEPSEEK_BASE_URL：API网关地址（https://openrouter.ai/api/v1）。")
    add_body("3. DEEPSEEK_MODEL：使用的模型标识（deepseek/deepseek-chat-v3-0324）。")
    add_body("4. TUSHARE_TOKEN：Tushare Pro数据接口令牌。")

    # ── 八、后续优化方向 ──
    add_h1("八、后续优化方向")

    add_body(
        "1. 流式输出：将当前的同步请求-响应模式升级为SSE流式输出，"
        "缩短用户感知的首字延迟。"
    )
    add_body(
        "2. 数据缓存层：引入Redis缓存高频查询结果（如当日行情、行业分类），"
        "降低Tushare API调用频次，提升响应速度。"
    )
    add_body(
        "3. 会话持久化：将内存中的会话数据迁移至数据库（如SQLite或PostgreSQL），"
        "支持服务重启后恢复历史对话。"
    )
    add_body(
        "4. 用户认证：增加登录认证机制，支持多用户独立会话管理。"
    )
    add_body(
        "5. 工具扩展：继续扩展工具集，覆盖可转债、基金、期权等更多品种。"
    )
    add_body(
        "6. 图表生成：集成图表渲染能力，对K线走势、资金流向等生成可视化图表。"
    )

    # ── Save ──
    out_path = os.path.join(os.path.dirname(__file__), "HQwenda技术设计文档.docx")
    doc.save(out_path)
    print(f"Document saved to: {out_path}")
    return out_path


if __name__ == "__main__":
    create_doc()
